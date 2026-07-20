import io
import re
import os
import json
import uuid as uuid_module
import logging
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from pydantic import BaseModel
from typing import Optional, List, Literal
from uuid import UUID
from app.core.config import settings
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from app.core.security import require_faculty
from app.db.supabase import get_supabase_admin

logger = logging.getLogger(__name__)
router = APIRouter()

QuestionType          = Literal["MCQ", "MSQ", "TRUE_FALSE", "SHORT_ANSWER", "LONG_ANSWER"]
Difficulty            = Literal["EASY", "MEDIUM", "HARD"]
ExtractedQuestionType = Literal["MCQ", "MSQ", "TRUE_FALSE", "SHORT_ANSWER"]

# ── Allowed image types for question image uploads ─────────────────────────────
ALLOWED_IMAGE_TYPES = {"image/jpeg", "image/png", "image/gif", "image/webp"}
MAX_IMAGE_SIZE_BYTES = 5 * 1024 * 1024  # 5 MB


# ── Ligature repair ────────────────────────────────────────────────────────────

# Some PDF font subsets encode "fi"/"fl"/etc. as a single ligature glyph.
# PyMuPDF generally decomposes these correctly, but occasionally emits the
# actual ligature codepoint instead (e.g. "ﬁ"). Expand any that slip through
# so downstream regex parsing and DB storage always see plain ASCII letters.
_LIGATURE_MAP = {
    "\ufb00": "ff", "\ufb01": "fi", "\ufb02": "fl",
    "\ufb03": "ffi", "\ufb04": "ffl", "\ufb05": "st", "\ufb06": "st",
}

# Some PDFs (commonly Google Forms "Print to PDF" exports) embed a subsetted
# font whose "fi" ligature glyph has a broken/missing ToUnicode CMap entry.
# Instead of decoding to the letters "f"+"i", PyMuPDF resolves it to whatever
# placeholder codepoint the broken table happens to point at — and this
# varies by font instance, so different quizzes concatenated into one PDF
# can each show a different artifact for the SAME missing ligature. We've
# observed it silently vanish (leaving e.g. "in nity" for "infinity") and,
# in other exports, render as U+00AC — the logical NOT sign, "¬" (e.g.
# "Path¬nding" for "Pathfinding").
#
# We can't safely strip every "¬" globally — it's also the legitimate
# logical-negation symbol and could appear in real question text (e.g.
# propositional-logic questions using "¬P"). But a real NOT sign never
# appears glued directly between two letters with no space or punctuation
# ("h¬n"); that specific pattern only happens when a ligature glyph has
# been mis-rendered mid-word, so it's a safe, unambiguous signal to repair.
_BROKEN_FI_GLYPH_RE = re.compile(r"(?<=[A-Za-z])¬(?=[A-Za-z])")


def _repair_ligatures(text: str) -> str:
    for lig, expansion in _LIGATURE_MAP.items():
        if lig in text:
            text = text.replace(lig, expansion)
    if "¬" in text:
        text = _BROKEN_FI_GLYPH_RE.sub("fi", text)
    return text


# ── Sanitize helper ───────────────────────────────────────────────────────────

def sanitize(text: str | None) -> str | None:
    """
    Strip characters that PostgreSQL cannot store in a `text` column, and
    repair any leftover ligature glyphs.

    PostgreSQL raises error code 22P05 ("unsupported Unicode escape sequence")
    when a string contains the NULL byte U+0000.  PDF text extraction can
    occasionally produce these from corrupt glyph tables.

    We also strip the Unicode replacement character U+FFFD which can appear
    for the same reason and is usually unwanted in question text, expand
    ligature glyphs (e.g. "ﬁ" -> "fi") that some PDF font subsets emit, and
    remove any leftover bold-detection sentinel markers as a final safety
    net before text reaches the database.
    """
    if text is None:
        return None
    text = text.replace("\x00", "").replace("\ufffd", "")
    text = _repair_ligatures(text)
    return _strip_bold_markers(text)


# ── Schemas ───────────────────────────────────────────────────────────────────

class OptionCreate(BaseModel):
    option_text: str
    is_correct: bool
    order_index: int

class TopicCreate(BaseModel):
    topic: str
    subject: Optional[str] = None
    chapter: Optional[str] = None

class QuestionCreate(BaseModel):
    course_id: Optional[UUID] = None
    question_type: QuestionType
    question_text: str
    marks: int
    negative_marks: int = 0
    difficulty: Difficulty
    options: Optional[List[OptionCreate]] = []
    topics: Optional[List[TopicCreate]] = []
    image_url: Optional[str] = None          # optional image attached to question

class OptionUpdate(BaseModel):
    # Present for an option that already exists in the DB (so we can update
    # it in place rather than delete+recreate, which would orphan any past
    # student answers that reference the option's id). Absent for a brand
    # new option being added during this edit.
    id: Optional[UUID] = None
    option_text: str
    is_correct: bool
    order_index: int

class QuestionUpdate(BaseModel):
    question_text: Optional[str] = None
    marks: Optional[int] = None
    negative_marks: Optional[int] = None
    difficulty: Optional[Difficulty] = None
    is_active: Optional[bool] = None
    image_url: Optional[str] = None          # allow updating image too
    # FIX: previously there was no way to update a question's options (e.g.
    # to fix a wrong "correct answer", reword a choice, or add/remove one)
    # once the question had been created — this field, plus the handling in
    # update_question() below, adds that.
    options: Optional[List[OptionUpdate]] = None

class ExtractedOption(BaseModel):
    text: str
    is_correct: bool

class ExtractedQuestion(BaseModel):
    id: str
    question_text: str
    question_type: ExtractedQuestionType
    options: List[ExtractedOption]
    marks: int
    difficulty: Difficulty
    confidence: int
    needs_review: bool
    approved: bool = False


# ── CRUD Endpoints ────────────────────────────────────────────────────────────

@router.get("/")
async def list_questions(
    course_id: Optional[UUID] = None,
    question_type: Optional[QuestionType] = None,
    difficulty: Optional[Difficulty] = None,
    is_active: Optional[bool] = True,
    _: dict = Depends(require_faculty),
):
    supabase = get_supabase_admin()
    query = supabase.table("questions").select(
        "*, courses(name, code), question_options(*), question_topics(*)"
    )
    if course_id:
        query = query.eq("course_id", str(course_id))
    if question_type:
        query = query.eq("question_type", question_type)
    if difficulty:
        query = query.eq("difficulty", difficulty)
    if is_active is not None:
        query = query.eq("is_active", is_active)
    return query.execute().data


@router.get("/{question_id}")
async def get_question(question_id: UUID, _: dict = Depends(require_faculty)):
    supabase = get_supabase_admin()
    result = (
        supabase.table("questions")
        .select("*, courses(name, code), question_options(*), question_topics(*)")
        .eq("id", str(question_id))
        .single()
        .execute()
    )
    if not result.data:
        raise HTTPException(status_code=404, detail="Question not found")
    return result.data


@router.post("/")
async def create_question(
    body: QuestionCreate,
    current_user: dict = Depends(require_faculty),
):
    supabase = get_supabase_admin()

    # FIX: sanitize() removes U+0000 null bytes and repairs ligature glyphs
    # that PDF extraction can embed in text, which PostgreSQL rejects with
    # error code 22P05.
    q_data = {
        "created_by":     current_user["user_id"],
        "course_id":      str(body.course_id) if body.course_id else None,
        "question_type":  body.question_type,
        "question_text":  sanitize(body.question_text),   # ← FIX
        "marks":          body.marks,
        "negative_marks": body.negative_marks,
        "difficulty":     body.difficulty,
        "is_active":      True,
        "image_url":      body.image_url,
    }
    q_result = supabase.table("questions").insert(q_data).execute()
    qid = q_result.data[0]["id"]

    if body.options:
        supabase.table("question_options").insert([
            {
                "question_id": qid,
                "option_text": sanitize(o.option_text),   # ← FIX
                "is_correct":  o.is_correct,
                "order_index": o.order_index,
            }
            for o in body.options
        ]).execute()

    if body.topics:
        supabase.table("question_topics").insert([
            {
                "question_id": qid,
                "topic":       sanitize(t.topic),         # ← FIX
                "subject":     sanitize(t.subject),       # ← FIX
                "chapter":     sanitize(t.chapter),       # ← FIX
            }
            for t in body.topics
        ]).execute()

    return {"message": "Question created", "question_id": qid}


# ── Image Upload Endpoint ─────────────────────────────────────────────────────
@router.post("/{question_id}/upload-image")
async def upload_question_image(
    question_id: UUID,
    file: UploadFile = File(...),
    current_user: dict = Depends(require_faculty),
):
    """
    Upload an image for a question. Stores in Supabase Storage bucket
    'question-images', updates the question's image_url column, and returns
    the public URL.

    Bucket setup (run once in Supabase dashboard):
        CREATE BUCKET question-images WITH (public = true);
    Or via SQL:
        INSERT INTO storage.buckets (id, name, public)
        VALUES ('question-images', 'question-images', true);
    """
    supabase = get_supabase_admin()

    # ── Validate question exists and belongs to this faculty ──────────────────
    result = (
        supabase.table("questions")
        .select("id, created_by")
        .eq("id", str(question_id))
        .single()
        .execute()
    )
    if not result.data:
        raise HTTPException(status_code=404, detail="Question not found")
    if result.data["created_by"] != current_user["user_id"]:
        raise HTTPException(status_code=403, detail="You can only upload images for your own questions")

    # ── Validate file type ────────────────────────────────────────────────────
    content_type = file.content_type or ""
    if content_type not in ALLOWED_IMAGE_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"Only JPEG, PNG, GIF, and WebP images are allowed. Got: {content_type}",
        )

    # ── Read and validate file size ───────────────────────────────────────────
    content = await file.read()
    if len(content) > MAX_IMAGE_SIZE_BYTES:
        raise HTTPException(
            status_code=400,
            detail="Image too large. Maximum size is 5 MB.",
        )

    # ── Build a unique storage path ───────────────────────────────────────────
    ext = content_type.split("/")[-1]           # e.g. "jpeg", "png"
    storage_path = f"{question_id}/{uuid_module.uuid4()}.{ext}"

    # ── Upload to Supabase Storage ────────────────────────────────────────────
    try:
        supabase.storage.from_("question-images").upload(
            path=storage_path,
            file=content,
            file_options={"content-type": content_type, "upsert": "true"},
        )
    except Exception as exc:
        logger.error("Supabase Storage upload failed: %s", exc, exc_info=True)
        raise HTTPException(status_code=500, detail="Image upload failed. Please try again.")

    # ── Build public URL ──────────────────────────────────────────────────────
    # Supabase public bucket URL pattern:
    #   {SUPABASE_URL}/storage/v1/object/public/{bucket}/{path}
    supabase_url = settings.SUPABASE_URL.rstrip("/")
    public_url = f"{supabase_url}/storage/v1/object/public/question-images/{storage_path}"

    # ── Persist URL on the question row ──────────────────────────────────────
    supabase.table("questions").update({"image_url": public_url}).eq("id", str(question_id)).execute()

    logger.info("Uploaded image for question %s → %s", question_id, public_url)
    return {"image_url": public_url}


# ── Delete Image Endpoint ─────────────────────────────────────────────────────
@router.delete("/{question_id}/image")
async def delete_question_image(
    question_id: UUID,
    current_user: dict = Depends(require_faculty),
):
    """Remove the image from a question (clears image_url; does NOT delete from storage)."""
    supabase = get_supabase_admin()
    result = (
        supabase.table("questions")
        .select("id, created_by, image_url")
        .eq("id", str(question_id))
        .single()
        .execute()
    )
    if not result.data:
        raise HTTPException(status_code=404, detail="Question not found")
    if result.data["created_by"] != current_user["user_id"]:
        raise HTTPException(status_code=403, detail="Forbidden")

    supabase.table("questions").update({"image_url": None}).eq("id", str(question_id)).execute()
    return {"message": "Image removed"}


def _replace_question_options(supabase, question_id: UUID, options: List[OptionUpdate]) -> None:
    """
    Persist an edited option list for an existing question.

    Options that carry an `id` are updated in place (so past student
    answers that reference that option's id via selected_option_id /
    selected_option_ids stay valid). Options without an `id` are newly
    added rows. Any existing option not present in the submitted list
    is removed.
    """
    existing = (
        supabase.table("question_options")
        .select("id")
        .eq("question_id", str(question_id))
        .execute()
    ).data or []
    existing_ids = {row["id"] for row in existing}
    submitted_ids = {str(o.id) for o in options if o.id}

    to_delete = existing_ids - submitted_ids
    if to_delete:
        supabase.table("question_options").delete().in_("id", list(to_delete)).execute()

    for o in options:
        if o.id and str(o.id) in existing_ids:
            supabase.table("question_options").update({
                "option_text": sanitize(o.option_text),
                "is_correct":  o.is_correct,
                "order_index": o.order_index,
            }).eq("id", str(o.id)).execute()
        else:
            supabase.table("question_options").insert({
                "question_id":  str(question_id),
                "option_text":  sanitize(o.option_text),
                "is_correct":   o.is_correct,
                "order_index":  o.order_index,
            }).execute()


@router.patch("/{question_id}")
async def update_question(
    question_id: UUID,
    body: QuestionUpdate,
    _: dict = Depends(require_faculty),
):
    supabase = get_supabase_admin()
    data = body.model_dump(exclude_none=True, exclude={"options"})
    if "question_text" in data:
        data["question_text"] = sanitize(data["question_text"])   # ← FIX: same null-byte/ligature safety net as create_question

    if not data and body.options is None:
        raise HTTPException(status_code=400, detail="No fields to update")

    result = None
    if data:
        result = supabase.table("questions").update(data).eq("id", str(question_id)).execute()

    if body.options is not None:
        _replace_question_options(supabase, question_id, body.options)

    if result and result.data:
        return result.data[0]
    # Only options changed (no `questions` row fields in this request) —
    # fetch and return the current row so callers still get a full object back.
    fetched = supabase.table("questions").select("*").eq("id", str(question_id)).single().execute()
    return fetched.data


@router.delete("/{question_id}")
async def soft_delete_question(question_id: UUID, _: dict = Depends(require_faculty)):
    supabase = get_supabase_admin()
    supabase.table("questions").update({"is_active": False}).eq("id", str(question_id)).execute()
    return {"message": "Question deactivated"}


# ── Text Extraction ───────────────────────────────────────────────────────────

# Bold text is wrapped in these sentinel markers during extraction so that
# downstream parsers can detect bold-marked correct answers (a very common
# convention in exported question banks: the correct option is simply bold,
# with no other marker like "*" or "(correct)"). Markers are stripped from
# question/option text before it is ever stored or shown to the user.
BOLD_OPEN  = "\u2060§B§\u2060"
BOLD_CLOSE = "\u2060§/B§\u2060"
BOLD_RUN_RE  = re.compile(re.escape(BOLD_OPEN) + r"(.*?)" + re.escape(BOLD_CLOSE), re.DOTALL)
BOLD_TAG_RE  = re.compile(re.escape(BOLD_OPEN) + "|" + re.escape(BOLD_CLOSE))

# Same idea, but for COLORED text — another very common way exam papers
# mark the correct answer (e.g. the correct option printed in red), often
# with no bold styling and no textual marker like "*" at all.
COLOR_OPEN  = "\u2060§C§\u2060"
COLOR_CLOSE = "\u2060§/C§\u2060"
COLOR_RUN_RE = re.compile(re.escape(COLOR_OPEN) + r"(.*?)" + re.escape(COLOR_CLOSE), re.DOTALL)
COLOR_TAG_RE = re.compile(re.escape(COLOR_OPEN) + "|" + re.escape(COLOR_CLOSE))

# Combined pattern used whenever we just need to strip ALL style markers to
# get back to plain, storable text (final question/option text should never
# contain either kind of sentinel marker).
ALL_STYLE_TAG_RE = re.compile(
    "|".join(re.escape(t) for t in (BOLD_OPEN, BOLD_CLOSE, COLOR_OPEN, COLOR_CLOSE))
)

BOLD_FONT_FLAG = 1 << 4  # PyMuPDF span flag bit for bold


def _span_is_bold(span: dict) -> bool:
    if span.get("flags", 0) & BOLD_FONT_FLAG:
        return True
    font_name = (span.get("font") or "").lower()
    return "bold" in font_name or "black" in font_name or "heavy" in font_name


def _span_is_colored(span: dict) -> bool:
    """
    True if this span's text color is anything other than plain black.
    PyMuPDF reports color as a packed int sRGB value; body text is almost
    always exactly 0 (pure black), so any non-zero value is a strong signal
    the author deliberately colored this text — commonly used in printed
    exam papers to mark the correct option (e.g. in red) without any other
    textual indicator.
    """
    color = span.get("color")
    return color is not None and color != 0


def _extract_pdf_text(content: bytes) -> str:
    """
    Extract text from a PDF using PyMuPDF (fitz), preserving bold AND color
    styling as inline sentinel markers (BOLD_OPEN/CLOSE, COLOR_OPEN/CLOSE)
    so that a bold-only or colored-only answer convention can be detected
    downstream.

    PyMuPDF decomposes ligature glyphs (e.g. the single "fi" glyph some PDF
    font subsets use) back into their constituent letters far more reliably
    than pdfplumber, which tends to silently drop them (producing corrupted
    text like "in nity" instead of "infinity"). Any ligature codepoints that
    still slip through are repaired below as a safety net.

    CRITICAL: blocks and lines are explicitly re-sorted by their vertical
    (then horizontal) position on the page before being joined into text.
    Google Forms' "print to PDF" export writes its underlying content
    stream out of visual order — a question's title text box is frequently
    emitted into the stream AFTER that question's answer-option blocks,
    even though it renders visually above them. Reading the raw
    content-stream order (as plain `page.get_text("text")` does) silently
    misattributes titles to the wrong question, which is why some
    questions end up with an empty/garbled stem, or a stem borrowed from a
    neighboring question. Sorting by actual page position restores true
    reading order and fixes this at the source.

    CRITICAL #2: lines WITHIN the same block are joined with a single
    SPACE, not a newline — a "block" from PyMuPDF corresponds to one
    paragraph/label (e.g. one Google Forms answer option, or one question
    title), and multiple "lines" inside it are just word-wrap of that same
    paragraph. Only DIFFERENT blocks are joined with newlines. Without
    this distinction, a long question stem or answer option that happens
    to wrap onto a second physical line gets emitted as if it were a
    brand-new option/question — this was previously worked around with a
    fragile "starts with lowercase = continuation" heuristic downstream,
    which itself broke whenever a real option legitimately started with a
    lowercase letter (a common style, e.g. "determining good values for
    all the weights...").
    """
    parts: list[str] = []
    with fitz.open(stream=content, filetype="pdf") as pdf:
        for page in pdf:
            page_dict = page.get_text("dict")
            blocks = page_dict.get("blocks", [])

            # Sort blocks top-to-bottom, then left-to-right, by their bbox —
            # this is the fix for Google Forms' scrambled content-stream order.
            def _block_key(b):
                bbox = b.get("bbox", (0, 0, 0, 0))
                return (round(bbox[1], 1), round(bbox[0], 1))
            blocks_sorted = sorted(blocks, key=_block_key)

            page_lines: list[str] = []
            for block in blocks_sorted:
                lines = block.get("lines", [])
                lines_sorted = sorted(
                    lines,
                    key=lambda l: (round(l.get("bbox", (0, 0, 0, 0))[1], 1),
                                   round(l.get("bbox", (0, 0, 0, 0))[0], 1)),
                )
                block_line_texts: list[str] = []
                for line in lines_sorted:
                    line_text = ""
                    for span in line.get("spans", []):
                        text = span.get("text", "")
                        if not text:
                            continue
                        is_bold    = _span_is_bold(span) and text.strip()
                        is_colored = _span_is_colored(span) and text.strip()
                        opened = ""
                        closed = ""
                        if is_bold:
                            opened += BOLD_OPEN
                            closed = BOLD_CLOSE + closed
                        if is_colored:
                            opened += COLOR_OPEN
                            closed = COLOR_CLOSE + closed
                        line_text += f"{opened}{text}{closed}" if (opened or closed) else text
                    if line_text.strip():
                        block_line_texts.append(line_text.strip())
                if block_line_texts:
                    joined_block = " ".join(block_line_texts)
                    # Skip standalone page-number blocks (e.g. a lone "1" or
                    # "2" printed as a page footer/header). Left as-is, these
                    # can leak into an option's text as stray trailing
                    # digits whenever a question's options happen to span a
                    # page break, since the block-slicing logic that splits
                    # "A) .. B) .. C) .. D) .." options has no other way to
                    # know a page boundary (and the accompanying footer)
                    # sits between two options rather than within one.
                    if re.fullmatch(r"\d{1,3}", _strip_bold_markers(joined_block).strip()):
                        continue
                    # Join wrapped lines of the SAME paragraph with a space.
                    page_lines.append(joined_block)
            page_text = "\n".join(page_lines).strip()
            if page_text:
                parts.append(_repair_ligatures(page_text))
    return "\n\n".join(parts)


def _is_option_line_bold(raw_line: str, threshold: float = 0.7) -> bool:
    """
    True if a sufficiently large fraction of the visible (non-marker)
    characters in this option line are wrapped in bold sentinel markers.
    Guards against a single bold word inside an otherwise plain option
    counting as a "bold answer" signal.
    """
    visible = ALL_STYLE_TAG_RE.sub("", raw_line).strip()
    if not visible:
        return False
    bold_chars = sum(len(m) for m in BOLD_RUN_RE.findall(raw_line))
    return bold_chars >= threshold * len(visible)


def _is_option_line_colored(raw_line: str, threshold: float = 0.7) -> bool:
    """
    Same idea as _is_option_line_bold, but for non-black text color —
    catches papers that mark the correct answer purely by color (e.g. red
    text) with no bold styling and no textual marker at all.
    """
    visible = ALL_STYLE_TAG_RE.sub("", raw_line).strip()
    if not visible:
        return False
    colored_chars = sum(len(m) for m in COLOR_RUN_RE.findall(raw_line))
    return colored_chars >= threshold * len(visible)


def _strip_bold_markers(text: str) -> str:
    """
    Strips ALL style sentinel markers (both bold and color) — kept under
    its original name for backward compatibility with existing call sites,
    but covers both marker kinds since final stored text must never
    contain either.
    """
    return ALL_STYLE_TAG_RE.sub("", text)


def _extract_docx_text(content: bytes) -> str:
    """
    Extract text from a DOCX, preserving bold-run styling via the same
    sentinel markers used for PDFs, so a bold-only correct-answer
    convention is detected consistently across both file types.
    """
    doc = DocxDocument(io.BytesIO(content))
    parts: list[str] = []

    def _para_to_marked_text(para) -> str:
        out = ""
        for run in para.runs:
            text = run.text
            if not text:
                continue
            if run.bold and text.strip():
                out += f"{BOLD_OPEN}{text}{BOLD_CLOSE}"
            else:
                out += text
        return out if out else para.text

    for para in doc.paragraphs:
        t = _para_to_marked_text(para).strip()
        if t:
            parts.append(_repair_ligatures(t))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t = _para_to_marked_text(para).strip()
                    if t:
                        parts.append(_repair_ligatures(t))
    return "\n".join(parts)


# ── Regex Patterns ────────────────────────────────────────────────────────────

SKIP_LINE_RE = re.compile(
    r"(?i)^("
    r"mark\s+only\s+one\s+oval\.?"
    r"|mark\s+all\s+that\s+apply\.?"
    r"|required\s+question"
    r"|\*\s*indicates\s+required(\s+question)?"
    r"|this\s+(form|content)\s+(doesn'?t.*|is\s+neither.*)"
    r"|google\s*\.?\s*forms?"
    r"|dear\s+students,?"
    r"|the\s+quiz\s+will\s+be\s+open.*"
    r"|start\s+time\s*:.*"
    r"|stop\s+time\s*:.*"
    r"|all\s+the\s+best!?"
    r"|roll\s+no\.?\s*\*?"
    r"|name\s*\*?"
    r"|email\s*\*?"
    r"|once\s+submit.*"
    r"|you\s+have\s+to\s+(be|wait).*"
    r"|its?\s+responsibility.*"
    r"|there\s+will\s+be\s+no\s+make.*"
    r"|the\s+attendance\s+is\s+mandatory.*"
    r"|forms\s*$"
    # Course/quiz title lines (e.g. "Quiz-II Artificial Intelligence
    # July-Dec 2024"). When multiple Google Forms quizzes are concatenated
    # into a single PDF, each quiz's own title/header block can land on the
    # SAME page as the previous quiz's last question, so the page-level
    # header-page filter alone can't catch it — this line-level pattern
    # (any line mentioning "quiz" and containing a 4-digit year) does.
    r"|quiz[\s\-]*[ivx\d]*\b.*\d{4}\s*$"
    r")\s*$",
)

# Lines that mark the start of a NEW quiz's front-matter block. Used to
# force-close whatever question is currently being built (`current`) when
# a second (or third...) quiz's header appears mid-stream in a
# multi-quiz-per-PDF export, so trailing header/footer noise never gets
# appended onto the previous quiz's last question as fake options.
NEW_QUIZ_BOUNDARY_RE = re.compile(
    r"(?i)^(dear\s+students,?|quiz[\s\-]*[ivx\d]*\b.*\d{4})\s*$"
)

QUESTION_NUM_RE  = re.compile(r"^[ \t]*(\d{1,3})\.\s*(.*)?$")
MARKS_RE         = re.compile(r"(?i)\b(\d+)\s*(?:points?|marks?)\b")
TRUE_FALSE_RE    = re.compile(r"(?i)\btrue\s*/?\s*false\b")
OPTION_LABEL_RE  = re.compile(r"^[ \t]*[\(\[]?([A-Da-d])[\).\:\]][ \t]+(.+)$")
CORRECT_MARKER_RE = re.compile(
    r"(?i)[ \t]*(?:[*\u2713\u2714]+[ \t]*"
    r"|\(correct\)|\[correct\]"
    r"|\(ans(?:wer)?\)|\[ans(?:wer)?\])[ \t]*$"
)
ANSWER_KEY_HEADER_RE = re.compile(
    r"(?im)^[ \t]*(answer[\s_-]*key|answers?|solutions?|key)[ \t]*[:\-]?[ \t]*$"
)
ANSWER_KEY_ENTRY_RE = re.compile(
    r"(?im)^[ \t]*(\d{1,3})[ \t]*[.):\-][ \t]*"
    r"([A-Da-d](?:[ \t]*[,\/][ \t]*[A-Da-d])*|true|false)[ \t]*$"
)


# ── Format Detection ──────────────────────────────────────────────────────────

def _detect_format(text: str) -> str:
    stripped = _strip_bold_markers(text)
    lower = stripped.lower()

    gform_hits = sum(1 for phrase in [
        "mark only one oval",
        "mark all that apply",
        "this form doesn",
        "google forms",
        "indicates required question",
    ] if phrase in lower)
    if gform_hits >= 2:
        return "google_forms"

    # Printed university/institute exam papers: numbered "Q.1", "Q.2", ...
    # and/or a "Response Table" answer-key section. Checked with a fairly
    # high question-count threshold so a stray "Q.1" reference elsewhere
    # in a standard-format doc doesn't misfire.
    q_dot_hits = len(re.findall(r"(?i)\bQ\.\s?\d{1,3}\b", stripped))
    has_response_table = bool(re.search(r"(?im)^\s*response\s+table\s*$", stripped))
    if q_dot_hits >= 3 or (q_dot_hits >= 1 and has_response_table):
        return "university_qa"

    return "standard"


def _build_stripped_offset_map(raw: str) -> tuple[str, list[int]]:
    """
    Strip all style sentinel markers (bold/color) out of `raw`, returning
    both the stripped text and a per-character index map back into the
    original `raw` string. This lets us run ordinary structural regexes
    (question boundaries, answer-key headers, etc.) against clean text
    without losing track of where bold/color styling was, so we can still
    slice the ORIGINAL marker-bearing text for style-based correct-answer
    detection afterwards. `index_map[i]` gives the index in `raw`
    corresponding to `stripped[i]`.
    """
    out_chars: list[str] = []
    index_map: list[int] = []
    i = 0
    n = len(raw)
    tags = (BOLD_OPEN, BOLD_CLOSE, COLOR_OPEN, COLOR_CLOSE)
    while i < n:
        matched = False
        for tag in tags:
            if raw.startswith(tag, i):
                i += len(tag)
                matched = True
                break
        if matched:
            continue
        out_chars.append(raw[i])
        index_map.append(i)
        i += 1
    return "".join(out_chars), index_map


def _raw_index(idx_map: list[int], stripped_pos: int, raw_len: int) -> int:
    """Map a position in stripped text back to the original raw text."""
    if stripped_pos < len(idx_map):
        return idx_map[stripped_pos]
    return raw_len


def _is_header_page(page_text: str) -> bool:
    stripped_page = _strip_bold_markers(page_text)
    lines = [l.strip() for l in stripped_page.splitlines() if l.strip()]
    if any(QUESTION_NUM_RE.match(l) for l in lines):
        return False
    keywords = [
        "quiz", "dear students", "roll no", "name", "email",
        "start time", "stop time", "all the best", "open only for",
        "auto-submission", "10 mins",
    ]
    hits = sum(1 for kw in keywords if any(kw in l.lower() for l in lines))
    return hits >= 3


# ── Google Forms Parser ───────────────────────────────────────────────────────

def _parse_google_forms(full_text: str) -> list[dict]:
    """
    Bold-marker-aware Google Forms parser.

    Google Forms PDFs don't normally export an answer key, but some faculty
    manually re-export or edit the PDF to bold the correct option before
    sharing it as an answer key. We treat a (near-)entirely-bold option line
    as a correct-answer signal, on top of the existing question/option
    structural parsing — all of which now runs against a marker-stripped
    VIEW of each line, with the raw (marker-bearing) line kept alongside
    purely for the bold-ratio check.
    """
    pages = full_text.split("\n\n")
    content_lines: list[str] = []
    for page in pages:
        if _is_header_page(page):
            logger.debug("Skipping header page")
        else:
            content_lines.extend(page.splitlines())

    STANDALONE_ASTERISK_RE = re.compile(r"^\*+$")
    cleaned: list[str] = []  # list of (raw_line, stripped_line)
    for raw in content_lines:
        raw_line = raw.strip()
        stripped_line = _strip_bold_markers(raw_line).strip()
        if not stripped_line or SKIP_LINE_RE.match(stripped_line):
            continue
        if STANDALONE_ASTERISK_RE.match(stripped_line):
            continue
        cleaned.append((raw_line, stripped_line))

    questions: list[dict] = []
    current: dict | None = None
    awaiting_question_text = False

    for raw_line, line in cleaned:
        m = QUESTION_NUM_RE.match(line)
        if m:
            if current and current.get("question_text"):
                questions.append(current)

            number = int(m.group(1))
            rest   = (m.group(2) or "").strip()

            marks = 1
            mm = MARKS_RE.search(rest)
            if mm:
                marks = int(mm.group(1))
                rest  = MARKS_RE.sub("", rest).strip()

            rest = re.sub(r"\s*\*+\s*$", "", rest).strip()

            current = {
                "number":        number,
                "question_text": rest,
                "options":       [],
                "marks":         marks,
                "has_correct":   False,
                "ai_answered":   False,
            }
            awaiting_question_text = not bool(rest)
            continue

        if current is None:
            continue

        if awaiting_question_text and not current["question_text"]:
            q_text = MARKS_RE.sub("", line).strip()
            q_text = re.sub(r"\s*\*+\s*$", "", q_text).strip()
            if not q_text:
                continue
            current["question_text"] = q_text
            awaiting_question_text = False
            continue

        opt_text = MARKS_RE.sub("", line).strip()
        opt_text = re.sub(r"\s*\*+\s*$", "", opt_text).strip()
        if not opt_text:
            continue

        is_corr_bold = _is_option_line_bold(raw_line)

        # NOTE: no lowercase-starting-line merge here. Word-wrapped
        # continuations of the same paragraph are already reassembled by
        # the extractor (_extract_pdf_text joins lines within the same
        # PDF block with a space before this parser ever sees them), so
        # each `line` reaching this point is a genuinely distinct option.
        # A "starts with lowercase = continuation" heuristic used to live
        # here, but it incorrectly merged separate options that legitimately
        # start with a lowercase word (a common style, e.g. "determining
        # good values for...", "calculate the shortest path...").
        current["options"].append({"text": opt_text, "is_correct": is_corr_bold})

    if current and current.get("question_text"):
        questions.append(current)

    for q in questions:
        q["has_correct"] = any(o["is_correct"] for o in q["options"])

    logger.info("Google Forms parser found %d questions", len(questions))
    return questions


# ── University / Printed Exam Parser ──────────────────────────────────────────
#
# Handles structured, printed-exam-style PDFs (e.g. institute quiz papers)
# using "Q.N ..." numbering, options given inline as "A) .. B) .. C) .. D) .."
# (either all on one line or one per line), fill-in-the-blank questions (no
# options, "____" blanks in the stem), and — when present — a printed answer
# key: either a compact two-row "A1 A2 A3 ... / C C C ..." table, or explicit
# per-question written answers for fill-in-the-blank questions.
#
# Bold or colored option text (a very common way these papers visually mark
# the correct answer — e.g. printing it in red) is treated as a
# correct-answer signal directly, on top of anything found in a printed
# answer key.

OPTION_INLINE_RE = re.compile(r"(?<![A-Za-z0-9])([A-D])\)\s*")
BLANK_RE = re.compile(r"_{3,}")
ANSWER_TABLE_HEADER_RE = re.compile(r"(?i)\bA(\d{1,3})\b")
ANSWER_TABLE_VALUE_LINE_RE = re.compile(r"(?i)^\s*(?:[A-D]\s*)+$")
FILL_BLANK_ANSWER_START_RE = re.compile(r"(?m)(?:^|\s)(\d{1,3})\.\s+")


def _split_options_from_block(block_raw: str) -> tuple[str, list[dict]]:
    """
    Given the raw (style-marker-bearing) text of one question block —
    everything between this "Q.N" and the next — split off the question
    stem from its answer options. Options may appear all on one line
    ("A) .. B) .. C) .. D) ..") or one per line; both are handled the same
    way by scanning for "A)"/"B)"/"C)"/"D)" markers anywhere in the block
    rather than assuming line boundaries. Returns (stem_raw, options),
    where options is empty if fewer than 2 option markers were found (a
    strong signal this is a fill-in-the-blank / short-answer question
    rather than an MCQ).
    """
    block_stripped, block_idx_map = _build_stripped_offset_map(block_raw)
    matches = list(OPTION_INLINE_RE.finditer(block_stripped))
    if len(matches) < 2:
        return block_raw, []

    stem_end_raw = _raw_index(block_idx_map, matches[0].start(), len(block_raw))
    stem_raw = block_raw[:stem_end_raw]

    opts: list[dict] = []
    for i, m in enumerate(matches):
        letter = m.group(1).upper()
        start_s = m.end()
        end_s = matches[i + 1].start() if i + 1 < len(matches) else len(block_stripped)
        start_raw = _raw_index(block_idx_map, start_s, len(block_raw))
        end_raw   = _raw_index(block_idx_map, end_s, len(block_raw))
        opts.append({"letter": letter, "raw_text": block_raw[start_raw:end_raw]})
    return stem_raw, opts


def _apply_university_answer_key(key_raw: str, questions: list[dict]) -> None:
    """
    Parse the trailing answer-key section (everything from "Response Table"
    onward) and apply it to the already-parsed `questions` list in place.

    Two answer formats are handled:

    1. Explicit written answers for fill-in-the-blank questions, e.g.
       "4. Collapse to a single layer______." / "6. Forget and output" —
       matched by question number and assigned as free-text answers.

    2. A compact two-row MCQ table:
           A1 A2 A3 A4 A5 A6 A7 A8 A9 A10
           C  C  C  C  B  C  B
       The header row lists EVERY question number in that block, including
       fill-in-the-blank ones, but the value row only has one letter per
       actual MCQ question, in order — so the header is filtered down to
       just the MCQ question numbers (using the question types we already
       determined while parsing the body) before zipping it with the
       values, otherwise every answer from that point on would be
       misaligned.
    """
    key_stripped = _strip_bold_markers(key_raw)
    q_by_number = {q["number"]: q for q in questions}

    # ── Fill-in-the-blank explicit written answers ──
    # Bounded by whichever comes next: another "N." fill-blank entry, OR
    # the start of an "A1 A2 A3 ..." table header line — without the
    # latter, the LAST fill-blank answer in a block would otherwise
    # swallow an entire following MCQ answer table (nothing else marks
    # where its free-text answer is supposed to end).
    fillblank_marks = [
        (m.start(), m.end(), int(m.group(1))) for m in FILL_BLANK_ANSWER_START_RE.finditer(key_stripped)
    ]
    table_header_marks = [
        (hm.start(), hm.end(), None)
        for hm in re.finditer(r"(?m)^\s*(?:A\d{1,3}\s*){2,}$", key_stripped)
    ]
    combined = sorted(
        [(*m, "fillblank") for m in fillblank_marks] + [(*m, "header") for m in table_header_marks],
        key=lambda t: t[0],
    )
    for idx, (start, end, num, kind) in enumerate(combined):
        if kind != "fillblank":
            continue
        q = q_by_number.get(num)
        if not (q and q.get("is_short_answer")):
            continue
        text_end = combined[idx + 1][0] if idx + 1 < len(combined) else len(key_stripped)
        text = key_stripped[end:text_end].strip()
        text = BLANK_RE.sub("", text).strip().rstrip(".").strip()
        if text:
            q["short_answer"] = text
            q["has_correct"] = True

    # ── Tabular MCQ answer key ──
    lines = [l.strip() for l in key_stripped.splitlines() if l.strip()]
    i = 0
    while i < len(lines):
        # Sorted numerically rather than trusting extraction order — a
        # header line's own tokens can occasionally come out of order if
        # the PDF's underlying block positions are slightly irregular
        # (harmless for detection, but zipping must follow the answers'
        # true left-to-right / ascending-question-number order).
        header_nums = sorted(int(n) for n in ANSWER_TABLE_HEADER_RE.findall(lines[i]))
        if header_nums and i + 1 < len(lines) and ANSWER_TABLE_VALUE_LINE_RE.match(lines[i + 1]):
            values = re.findall(r"[A-D]", lines[i + 1].upper())
            mcq_nums = [
                n for n in header_nums
                if n in q_by_number and not q_by_number[n].get("is_short_answer")
            ]
            for num, letter in zip(mcq_nums, values):
                q = q_by_number[num]
                for opt in q["options"]:
                    opt["is_correct"] = (opt["letter"] == letter)
                q["has_correct"] = any(o["is_correct"] for o in q["options"])
            i += 2
            continue
        i += 1


def _parse_university_qa(full_text: str) -> list[dict]:
    stripped, idx_map = _build_stripped_offset_map(full_text)

    # Split off the trailing answer-key section ("Response Table" onward)
    # so it never gets mistaken for question content. Matched as a
    # standalone heading line (not just any mention of the phrase — it can
    # also appear inside the instructions, e.g. "Answers written in the
    # Response table will be checked only.").
    key_m = re.search(r"(?im)^\s*response\s+table\s*$", stripped)
    if key_m:
        split_raw = _raw_index(idx_map, key_m.start(), len(full_text))
        body_raw, key_raw = full_text[:split_raw], full_text[split_raw:]
    else:
        body_raw, key_raw = full_text, ""

    body_stripped, body_idx_map = _build_stripped_offset_map(body_raw)

    q_starts = list(re.finditer(r"(?i)Q\.\s*(\d{1,3})\b\.?\s*", body_stripped))
    if not q_starts:
        return []

    questions: list[dict] = []
    for i, m in enumerate(q_starts):
        number = int(m.group(1))
        start_s = m.end()
        end_s = q_starts[i + 1].start() if i + 1 < len(q_starts) else len(body_stripped)
        start_raw = _raw_index(body_idx_map, start_s, len(body_raw))
        end_raw   = _raw_index(body_idx_map, end_s, len(body_raw))
        block_raw = body_raw[start_raw:end_raw]

        stem_raw, opt_matches = _split_options_from_block(block_raw)
        stem = _strip_bold_markers(stem_raw).strip()
        stem = MARKS_RE.sub("", stem).strip()
        if not stem:
            continue

        if not opt_matches:
            # No A)/B)/C)/D) markers found -> fill-in-the-blank / short answer.
            questions.append({
                "number":          number,
                "question_text":   stem,
                "options":         [],
                "marks":           1,
                "has_correct":     False,
                "ai_answered":     False,
                "is_short_answer": True,
                "short_answer":    None,
            })
            continue

        opts: list[dict] = []
        for om in opt_matches:
            raw_opt_text = om["raw_text"]
            visible_opt_text = _strip_bold_markers(raw_opt_text).strip()
            visible_opt_text = re.sub(r"\s{2,}", " ", visible_opt_text)
            is_marked = _is_option_line_bold(raw_opt_text) or _is_option_line_colored(raw_opt_text)
            opts.append({
                "letter":     om["letter"],
                "text":       visible_opt_text,
                "is_correct": is_marked,
            })

        questions.append({
            "number":          number,
            "question_text":   stem,
            "options":         opts,
            "marks":           1,
            "has_correct":     any(o["is_correct"] for o in opts),
            "ai_answered":     False,
            "is_short_answer": False,
        })

    if key_raw:
        _apply_university_answer_key(key_raw, questions)

    logger.info("University Q&A parser found %d questions", len(questions))
    return questions


# ── Standard MCQ Parser ───────────────────────────────────────────────────────

def _parse_standard(full_text: str) -> list[dict]:
    """
    Bold-marker-aware standard MCQ parser.

    Structural regexes (question-number boundaries, answer-key header/entries)
    are always matched against a marker-stripped VIEW of the text so bold
    styling around a number or heading can never break detection. Because
    that stripped view is character-for-character shorter than the raw text
    wherever bold runs occur, we build both views together with an explicit
    per-line pairing (raw line -> stripped line) rather than doing index
    arithmetic across two differently-lengthed strings.
    """
    stripped_full = _strip_bold_markers(full_text)

    # ── Answer key extraction (structural — always on stripped text) ──
    answer_key: dict[int, list[str]] = {}
    hm = ANSWER_KEY_HEADER_RE.search(stripped_full)
    if hm:
        # Cut both the raw and stripped text at the header boundary. Since
        # everything before the header contains no answer-key content, we
        # simply re-derive the raw cutoff by taking the same number of
        # question-bearing lines rather than a raw character offset.
        stripped_before_header = stripped_full[: hm.start()]
        key_section             = stripped_full[hm.end():]

        n_body_lines = stripped_before_header.count("\n") + (1 if stripped_before_header else 0)
        raw_lines_all = full_text.splitlines()
        body_raw_lines = raw_lines_all[: min(n_body_lines, len(raw_lines_all))] if n_body_lines else raw_lines_all
        full_text = "\n".join(body_raw_lines)

        for entry in ANSWER_KEY_ENTRY_RE.finditer(key_section):
            q_num = int(entry.group(1))
            raw   = entry.group(2).strip().lower()
            if raw in ("true", "false"):
                answer_key[q_num] = [raw.upper()]
            else:
                answer_key[q_num] = [
                    p.strip().upper()
                    for p in re.split(r"[,\/]", raw) if p.strip()
                ]

    # ── Question boundaries (structural — matched on stripped text, then
    #    mapped back to raw lines by line index, which bold markers never
    #    change since they never introduce or remove newlines) ──
    raw_lines      = full_text.splitlines()
    stripped_lines = [_strip_bold_markers(l) for l in raw_lines]

    q_start_re = re.compile(r"^[ \t]*(\d{1,3})[.):\]]\s*(.*)$")
    q_line_idx: list[tuple[int, int]] = []  # (line_index, question_number)
    for idx, sline in enumerate(stripped_lines):
        m = q_start_re.match(sline)
        if m:
            q_line_idx.append((idx, int(m.group(1))))

    questions: list[dict] = []
    for i, (start_idx, number) in enumerate(q_line_idx):
        end_idx     = q_line_idx[i + 1][0] if i + 1 < len(q_line_idx) else len(raw_lines)
        block_lines = raw_lines[start_idx:end_idx]
        # Strip the leading "N." from just the first line before processing.
        first_stripped = stripped_lines[start_idx]
        m0 = q_start_re.match(first_stripped)
        # Rebuild the first line's remainder from the stripped match (the
        # question number prefix is discarded; it's never treated as an
        # option so losing its bold styling, if any, is harmless).
        block_lines = list(block_lines)
        block_lines[0] = m0.group(2) if m0 else block_lines[0]

        lines    = [l.strip() for l in block_lines if l.strip()]
        opts: list[dict] = []
        q_lines: list[str] = []
        seen_opt = False

        for ln in lines:
            # Match the option label directly on the raw (marker-bearing)
            # line — labels like "C)" are essentially never bold, so this
            # matches the same way it would on stripped text, and it lets us
            # isolate group(2) as the raw answer content (still carrying its
            # bold markers) separately from the label prefix.
            lm = OPTION_LABEL_RE.match(ln)
            if lm:
                seen_opt = True
                letter        = lm.group(1).upper()
                raw_content   = lm.group(2).strip()          # may contain bold markers
                visible_content = _strip_bold_markers(raw_content).strip()
                is_corr_marker = bool(CORRECT_MARKER_RE.search(visible_content))
                opt_text = CORRECT_MARKER_RE.sub("", visible_content).strip()
                # Bold-only convention: no explicit marker, but the answer
                # content itself (excluding the "C)" label) is (almost)
                # entirely bold-styled.
                is_corr_bold = (not is_corr_marker) and _is_option_line_bold(raw_content)
                opts.append({
                    "letter":     letter,
                    "text":       opt_text,
                    "is_correct": is_corr_marker or is_corr_bold,
                })
            elif not seen_opt:
                clean = MARKS_RE.sub("", _strip_bold_markers(ln)).strip()
                if clean:
                    q_lines.append(clean)

        q_text = " ".join(q_lines).strip()
        if not q_text or len(opts) < 2:
            continue

        marks = 1
        mm = MARKS_RE.search(_strip_bold_markers(" ".join(block_lines)))
        if mm:
            try: marks = int(mm.group(1))
            except: pass

        key = answer_key.get(number)
        if key and key not in (["TRUE"], ["FALSE"]):
            for opt in opts:
                opt["is_correct"] = opt.get("letter") in key

        has_correct = any(o["is_correct"] for o in opts)

        questions.append({
            "number":        number,
            "question_text": q_text,
            "options":       opts,
            "marks":         marks,
            "has_correct":   has_correct,
            "ai_answered":   False,
        })

    logger.info("Standard parser found %d questions", len(questions))
    return questions


# ── DeepSeek R1 (via OpenRouter) Answer Inference ─────────────────────────────

import requests
import time
import unicodedata

# Real batching: multiple questions go into ONE API call instead of one call
# per question. The previous DEEPSEEK_BATCH_SIZE = 1 meant a 20-question quiz
# could issue up to 20 x 9 = 180 request attempts (3 temperature/strictness
# configs x 3 retries each) with sleeps between every one — this is what
# produced 16-minute runtimes. Batching questions per call cuts a 20-question
# quiz down to ~4 requests total.
DEEPSEEK_BATCH_SIZE = 6
THINK_TAG_RE = re.compile(r"<think>.*?</think>", flags=re.DOTALL)
JSON_OBJECT_RE = re.compile(r"\{.*\}", re.DOTALL)


class _RateLimited(Exception):
    """Raised internally when the DeepSeek/OpenRouter endpoint returns 429."""
    pass


def _clean_text(text: str) -> str:
    text = unicodedata.normalize("NFKC", text)
    # Also strip null bytes, repair ligature glyphs, and remove any leftover
    # bold sentinel markers here as a second line of defence before text is
    # sent to the model (question_text/option text should already be clean
    # by this point, but this guards against any parser edge case).
    text = text.replace("\x00", "").replace("\ufffd", "")
    text = _repair_ligatures(text)
    text = _strip_bold_markers(text)
    return re.sub(r"\s+", " ", text).strip()


def _strip_reasoning(content: str) -> str:
    """
    Remove DeepSeek R1's <think>...</think> chain-of-thought block.
    Some providers only emit a closing '</think>' with no visible opener
    (the opening portion was streamed/consumed elsewhere), so also handle
    a dangling '</think>' by keeping only what follows it.
    """
    content = THINK_TAG_RE.sub("", content)
    if "</think>" in content:
        content = content.split("</think>", 1)[1]
    return content.strip()


def _build_batch_prompt(batch: list[dict]) -> str:
    lines: list[str] = []
    for i, q in enumerate(batch):
        q_text = _clean_text(q["question_text"])
        lines.append(f"Q{i}: {q_text}")
        for j, o in enumerate(q["options"]):
            lines.append(f"  {j}. {_clean_text(o['text'])}")
        lines.append("")
    return "\n".join(lines)


def _extract_json_object(content: str) -> dict | None:
    content = content.strip()
    content = re.sub(r"^```(?:json)?\s*|\s*```$", "", content, flags=re.MULTILINE).strip()
    try:
        return json.loads(content)
    except json.JSONDecodeError:
        pass
    m = JSON_OBJECT_RE.search(content)
    if not m:
        return None
    try:
        return json.loads(m.group(0))
    except json.JSONDecodeError:
        return None


def _call_deepseek_batch(api_key: str, batch: list[dict], temperature: float) -> dict[int, list[int]]:
    """
    Send an entire batch of questions in ONE request. The model returns a
    single JSON object mapping question index -> list of correct option
    index(es), which lets us resolve many questions per round-trip instead
    of one.
    """
    system_prompt = (
        "You are an expert academic tutor. You will be given several "
        "multiple-choice questions, each labeled Qn with its options "
        "numbered below it. For EACH question, determine the correct "
        "option index (or indices, if more than one option is correct).\n\n"
        "Respond with ONLY a single compact JSON object mapping each "
        "question's number (as a string key) to a list of its correct "
        "option index(es) — no explanation, no markdown, no code fences, "
        "nothing before or after the JSON. Always include an answer for "
        "every question, picking your best guess even if unsure.\n\n"
        "Example response for 3 questions:\n"
        '{"0": [1], "1": [0, 2], "2": [3]}'
    )
    user_content = _build_batch_prompt(batch)

    payload = {
        "model":       "openrouter/free",
        "temperature": temperature,
        # FIX: the "openrouter/free" router (see class comment above) can
        # land on ANY currently-available free model — not just DeepSeek R1
        # — and different providers format chain-of-thought differently.
        # Some wrap it in <think>...</think> (which _strip_reasoning
        # handles); others just write it as plain prose with no tags at
        # all ("We need to output a JSON object mapping..."), which is
        # exactly what showed up unparsed in production logs. Rather than
        # trying to detect every provider's reasoning style after the
        # fact, tell OpenRouter to keep it out of `content` in the first
        # place — this is documented as supported across all models.
        "reasoning": {"exclude": True},
        # FIX: 300 tokens/question was sized for "just emit the JSON",
        # not for a reasoning model's internal thinking pass — even with
        # reasoning.exclude above, some providers still consume part of
        # the same completion budget while thinking before the visible
        # answer appears. A batch of 5 questions was hitting this ceiling
        # and getting cut off before ever reaching the JSON. Floor of
        # 2500 plus real per-question headroom fixes that regardless of
        # which model answers.
        "max_tokens":  max(2500, 900 * len(batch) + 500),
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_content},
        ],
    }

    resp = requests.post(
        "https://openrouter.ai/api/v1/chat/completions",
        json=payload,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type":  "application/json",
            "HTTP-Referer":  settings.FRONTEND_URL,
            "X-Title":       "ExamPortal",
            "User-Agent":    "Mozilla/5.0 (compatible; ExamPortal/1.0)",
        },
        timeout=120,
    )

    if resp.status_code == 429:
        raise _RateLimited()
    resp.raise_for_status()

    body = resp.json()
    routed_model = body.get("model", "unknown")
    message = body["choices"][0]["message"]
    content = message.get("content", "") or ""
    content = _strip_reasoning(content)

    parsed = _extract_json_object(content)
    if not parsed:
        # Note which model the "openrouter/free" router actually picked —
        # useful for spotting a specific free model that's consistently
        # unreliable for this task — and whether it leaked reasoning into
        # a separate `reasoning` field despite reasoning.exclude, which
        # would point at a provider that doesn't honor that flag.
        leaked_reasoning = message.get("reasoning")
        logger.warning(
            "DeepSeek batch response was not valid JSON (model=%s, reasoning_leaked=%s): %r",
            routed_model, bool(leaked_reasoning), content[:300],
        )
        return {}

    result: dict[int, list[int]] = {}
    for k, v in parsed.items():
        try:
            idx = int(k)
        except (TypeError, ValueError):
            continue
        if isinstance(v, list):
            nums = [int(x) for x in v if str(x).lstrip("-").isdigit()]
        elif str(v).lstrip("-").isdigit():
            nums = [int(v)]
        else:
            continue
        result[idx] = nums
    return result


def _infer_answers_batch(api_key: str, batch: list[dict]) -> None:
    """
    Try to resolve every question in `batch` with as few requests as
    possible: one attempt at temperature 0, and — only for whatever
    questions are still unanswered — a single retry at a higher
    temperature. This replaces the old 3-config x 3-retry-per-question
    scheme, which was the main source of the 16-minute runtime.
    """
    remaining = list(enumerate(batch))  # (local_index, question_dict)

    for temperature in (0, 0.5):
        if not remaining:
            return

        sub_batch = [q for _, q in remaining]
        backoff = 1.0
        mapping: dict[int, list[int]] = {}
        for _retry in range(2):  # up to 2 tries per temperature on 429
            try:
                mapping = _call_deepseek_batch(api_key, sub_batch, temperature=temperature)
                break
            except _RateLimited:
                logger.info("DeepSeek batch rate-limited, backing off %.1fs", backoff)
                time.sleep(backoff)
                backoff *= 2
            except (requests.exceptions.HTTPError, requests.exceptions.Timeout, json.JSONDecodeError) as e:
                logger.warning("DeepSeek batch request failed: %s", e)
                break

        still_remaining = []
        for pos, (local_idx, q) in enumerate(remaining):
            idxs = mapping.get(pos)
            valid = [x for x in (idxs or []) if 0 <= x < len(q["options"])]
            if valid:
                for j, opt in enumerate(q["options"]):
                    opt["is_correct"] = j in valid
                q["has_correct"] = True
                q["ai_answered"] = True
            else:
                still_remaining.append((local_idx, q))
        remaining = still_remaining
        time.sleep(0.3)

    if remaining:
        logger.warning("DeepSeek could not produce confident answers for %d question(s) in batch", len(remaining))


def _infer_answers_with_deepseek(raw_questions: list[dict]) -> list[dict]:
    api_key = settings.OPENROUTER_API_KEY.strip()
    if not api_key:
        # Pydantic requires OPENROUTER_API_KEY to be set at startup, but guard
        # against an empty string sneaking through via .env.
        logger.warning("OPENROUTER_API_KEY is blank — skipping AI answer inference")
        return raw_questions

    # Defense in depth: only ever send questions that (a) don't already
    # have a correct answer extracted from the PDF itself, and (b) are
    # MCQ/MSQ/TRUE_FALSE-shaped (have >= 2 options) — short-answer /
    # fill-in-the-blank questions are never AI-inferred. This mirrors the
    # filtering already done by the caller, but is enforced here too so
    # this function is safe to call directly without re-litigating the
    # "don't guess an answer that's already on the page" rule elsewhere.
    to_infer = [
        q for q in raw_questions
        if not q.get("has_correct") and not q.get("is_short_answer") and len(q.get("options", [])) >= 2
    ]
    if not to_infer:
        return raw_questions

    for start in range(0, len(to_infer), DEEPSEEK_BATCH_SIZE):
        batch = to_infer[start:start + DEEPSEEK_BATCH_SIZE]
        try:
            _infer_answers_batch(api_key, batch)
        except Exception as e:
            logger.error("DeepSeek inference failed (batch @%d): %s", start, e, exc_info=True)
        time.sleep(0.3)

    answered = sum(1 for q in to_infer if q.get("ai_answered") and q.get("has_correct"))
    logger.info("DeepSeek answered %d / %d questions that needed inference", answered, len(to_infer))
    return raw_questions


# ── Finalize to ExtractedQuestion ─────────────────────────────────────────────

def _finalize(raw_questions: list[dict]) -> list[dict]:
    results: list[dict] = []

    for q in raw_questions:
        q_text  = q.get("question_text", "").strip()
        marks   = q.get("marks", 1)
        number  = q.get("number", len(results) + 1)

        # ── Fill-in-the-blank / short-answer questions ──
        if q.get("is_short_answer"):
            if not q_text:
                logger.debug("Skipping Q%s — empty short-answer stem", number)
                continue

            short_answer = (q.get("short_answer") or "").strip()
            has_correct  = bool(short_answer)
            # The extracted answer text is stored as a single option with
            # is_correct=True, reusing the existing options structure
            # rather than requiring a schema change for free-text answers.
            final_opts = [{"text": short_answer, "is_correct": True}] if short_answer else []

            conf = 40
            if len(q_text) > 15: conf += 10
            if len(q_text) > 40: conf += 5
            if has_correct:      conf += 35
            conf = min(conf, 95)

            wc = len(q_text.split())
            difficulty = "EASY" if wc < 12 else ("MEDIUM" if wc < 30 else "HARD")

            results.append({
                "id":            f"ext-{len(results) + 1}",
                "question_text": q_text,
                "question_type": "SHORT_ANSWER",
                "options":       final_opts,
                "marks":         marks,
                "difficulty":    difficulty,
                "confidence":    conf,
                # Short-answer questions always need a human look — an
                # extracted written answer is a good starting point, but
                # free-text matching is inherently fuzzier than an MCQ
                # letter, so these are never treated as fully confident.
                "needs_review":  True,
                "approved":      False,
            })
            continue

        options = q.get("options", [])
        if not q_text or len(options) < 2:
            logger.debug("Skipping Q%s — insufficient text or options", number)
            continue

        opt_texts_lower = {o["text"].strip().lower() for o in options}
        is_tf = opt_texts_lower <= {"true", "false"} or bool(TRUE_FALSE_RE.search(q_text))

        if is_tf:
            question_type = "TRUE_FALSE"
            tf_correct: str | None = None
            for opt in options:
                if opt.get("is_correct"):
                    val = opt["text"].strip().capitalize()
                    if val in ("True", "False"):
                        tf_correct = val
                        break
            final_opts = [
                {"text": "True",  "is_correct": tf_correct == "True"},
                {"text": "False", "is_correct": tf_correct == "False"},
            ]
        else:
            correct_count = sum(1 for o in options if o.get("is_correct"))
            question_type = "MSQ" if correct_count > 1 else "MCQ"
            final_opts = [
                {"text": o["text"].strip(), "is_correct": o.get("is_correct", False)}
                for o in options
                if o.get("text", "").strip()
            ]

        has_correct = any(o["is_correct"] for o in final_opts)

        conf = 40
        if len(q_text) > 15:       conf += 10
        if len(q_text) > 40:       conf += 5
        if len(final_opts) >= 4:   conf += 10
        elif len(final_opts) >= 2: conf += 5
        if has_correct:            conf += 25
        if q.get("ai_answered"):   conf += 10
        conf = min(conf, 95)

        needs_review = not has_correct or conf < 70

        wc = len(q_text.split())
        difficulty = "EASY" if wc < 12 else ("MEDIUM" if wc < 30 else "HARD")

        results.append({
            "id":            f"ext-{len(results) + 1}",
            "question_text": q_text,
            "question_type": question_type,
            "options":       final_opts,
            "marks":         marks,
            "difficulty":    difficulty,
            "confidence":    conf,
            "needs_review":  needs_review,
            "approved":      False,
        })

    return results


# ── Extract Endpoint ──────────────────────────────────────────────────────────

@router.post("/extract", response_model=List[ExtractedQuestion])
async def extract_questions_from_file(
    file: UploadFile = File(...),
    _: dict = Depends(require_faculty),
):
    filename = (file.filename or "").lower()
    content  = await file.read()

    if not content:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    if filename.endswith(".pdf"):
        try:
            full_text = _extract_pdf_text(content)
        except Exception as exc:
            logger.error("PDF extraction failed: %s", exc, exc_info=True)
            raise HTTPException(
                status_code=400,
                detail=(
                    "Could not read this PDF. Ensure it is text-based "
                    "(exported from Word/Google Docs/LaTeX) and not password-protected."
                ),
            )
    elif filename.endswith(".docx"):
        try:
            full_text = _extract_docx_text(content)
        except Exception as exc:
            logger.error("DOCX extraction failed: %s", exc, exc_info=True)
            raise HTTPException(status_code=400, detail="Could not read this DOCX file.")
    else:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported.")

    if not full_text.strip():
        raise HTTPException(
            status_code=400,
            detail=(
                "No text could be extracted. "
                "Ensure the PDF is text-based, not a scanned image."
            ),
        )

    logger.info("Extracted %d chars from '%s'", len(full_text), filename)

    fmt = _detect_format(full_text)
    logger.info("Detected PDF format: %s", fmt)

    raw_questions = (
        _parse_google_forms(full_text) if fmt == "google_forms"
        else _parse_university_qa(full_text) if fmt == "university_qa"
        else _parse_standard(full_text)
    )

    if not raw_questions:
        raise HTTPException(
            status_code=422,
            detail=(
                "No questions could be identified. "
                "Ensure questions are numbered (1., 2., ...) and "
                "the PDF is a text-based export."
            ),
        )

    # Only MCQ/MSQ/TRUE_FALSE questions that don't already have a correct
    # answer extracted directly from the PDF (explicit marker, printed
    # answer key, or bold/color highlighting) are sent to the AI at all —
    # if the answer was already found in the document, the AI should never
    # be asked to guess it. Short-answer / fill-in-the-blank questions are
    # never sent to the (MCQ-oriented) inference prompt either; their
    # answers, when present, only ever come from an explicit written
    # answer key in the PDF itself.
    needs_ai = [
        q for q in raw_questions
        if not q.get("has_correct") and not q.get("is_short_answer") and len(q.get("options", [])) >= 2
    ]
    if needs_ai:
        logger.info(
            "%d / %d questions need AI answer inference (rest already answered from the PDF)",
            len(needs_ai), len(raw_questions),
        )
        raw_questions = _infer_answers_with_deepseek(raw_questions)
    else:
        logger.info("All questions already have an answer extracted from the PDF — skipping AI inference entirely")

    extracted = _finalize(raw_questions)

    if not extracted:
        raise HTTPException(
            status_code=422,
            detail=(
                "Questions were found but could not be structured. "
                "Each question needs at least 2 options."
            ),
        )

    logger.info("Returning %d extracted questions from '%s'", len(extracted), filename)
    return extracted